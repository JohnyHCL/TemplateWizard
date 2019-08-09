from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import docx, pyperclip, re

def finalTouch(tab):
    tab.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    tab.rows[1].cells[0].paragraphs[0].runs[0].font.underline = True
    tab.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
     
    for row in tab.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(12)

def paster():
    
    def parser(x):
        n = description.find(x)
        return (description[n:n+description[n:].find('\n')].split(':')[1] if n != -1 else -1).strip()
    
    doc = docx.Document('files\\template.docx')
    table = doc.tables[0]
    
    data = pyperclip.paste()
    data = data.split('/nextEl,')
    
    incNo = data[0] 
    incStatus = data[1] 
    incPrio = data[2] 
    summary = data[3]
    description = data[4]
    RG = data[5]
    startDate = data[6]
    latestDate = data[7]
    latestUpdate = data[8]
    descStartTime = parser('ISSUE START TIME:')
    desc = parser('ISSUE DESCRIPTION:')
    location = parser('LOCATION')
      
    table.cell(1,0).text = '\n' + 'P' + incPrio + ' ' + incNo + ' Incident Initial Notification' + '\n'
    table.cell(2,1).text = incNo
    table.cell(3,1).text = startDate 
    table.cell(4,1).text = 'P' + incPrio
    table.cell(4,3).text = incStatus
    table.cell(5,1).text = summary
    table.cell(6,1).text = (desc if desc != -1 else '')
    table.cell(10,1).text = location
    table.cell(11,1).text = 'Jan Sobczak'
    table.cell(11,3).text = RG
    table.cell(12,1).text = '1st SD Line'
    table.cell(13,1).text = latestDate + ' - ' + latestUpdate
    table.cell(14,1).text = ('30 minutes' if incPrio == '1' else 'Upon Resolution')
     
    finalTouch(table)
     
    doc.save('files\\output.docx')
    
def colors(x):
    
    doc = docx.Document('files\\color_change.docx')
    table = doc.tables[0]
    
    if x == '1':
        val = 'FF0000'
    elif x == '2':
        val = 'FFC000'
        table.cell(1,0).text = table.cell(1,0).text.replace('Initial', 'Update')
    elif x == '3':
        val = '00B050'
        table.cell(1,0).text = table.cell(1,0).text.replace('Initial', 'Final') 
    
    fill1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill2 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill3 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill4 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill5 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill6 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill7 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill8 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill9 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill10 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill11 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill12 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill13 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill14 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill15 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
    fill16 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        
    table.cell(2,0)._tc.get_or_add_tcPr().append(fill1)
    table.cell(3,0)._tc.get_or_add_tcPr().append(fill2)
    table.cell(4,0)._tc.get_or_add_tcPr().append(fill3)
    table.cell(4,2)._tc.get_or_add_tcPr().append(fill4)
    table.cell(5,0)._tc.get_or_add_tcPr().append(fill5)
    table.cell(6,0)._tc.get_or_add_tcPr().append(fill6)
    table.cell(7,0)._tc.get_or_add_tcPr().append(fill7)
    table.cell(7,2)._tc.get_or_add_tcPr().append(fill8)
    table.cell(8,0)._tc.get_or_add_tcPr().append(fill9)
    table.cell(9,0)._tc.get_or_add_tcPr().append(fill10)
    table.cell(10,0)._tc.get_or_add_tcPr().append(fill11)
    table.cell(11,0)._tc.get_or_add_tcPr().append(fill12)
    table.cell(11,2)._tc.get_or_add_tcPr().append(fill13)
    table.cell(12,0)._tc.get_or_add_tcPr().append(fill14)
    table.cell(13,0)._tc.get_or_add_tcPr().append(fill15)        
    table.cell(14,0)._tc.get_or_add_tcPr().append(fill16)
    
    finalTouch(table)
    
    doc.save('files\\output.docx')